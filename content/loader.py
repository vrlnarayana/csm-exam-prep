from pathlib import Path
import re

ROOT = Path(__file__).parent.parent


def get_day_content(day_number: int) -> dict:
    from content.days import DAYS
    day = next((d for d in DAYS if d["day"] == day_number), None)
    if day is None:
        return {"sections": [], "key_terms": {}, "title": f"Day {day_number}"}

    override_path = day.get("docx_override")
    if override_path:
        full_path = ROOT / override_path
        if full_path.exists():
            try:
                return _parse_study_docx(full_path, day)
            except Exception:
                pass

    return day


def get_day_questions(day_number: int) -> list:
    if day_number == 1:
        docx_path = ROOT / "Day1" / "CSM_100_Practice_Questions-Day 1.docx"
        if docx_path.exists():
            try:
                parsed = _parse_questions_docx(docx_path)
                if parsed:
                    return parsed
            except Exception:
                pass
    return _get_builtin_questions(day_number)


def _get_builtin_questions(day_number: int) -> list:
    from content.questions import QUESTIONS
    return [q for q in QUESTIONS if q["day"] == day_number]


def _parse_study_docx(path: Path, day_meta: dict) -> dict:
    """Parse the Day 1 study .docx.

    The file uses paragraph styles: Heading 2 / Heading 3 for section titles,
    List Paragraph for bullets, and None (empty style name) for body text.
    """
    from docx import Document
    doc = Document(str(path))
    sections = []
    current_section = None
    in_tip = False

    for para in doc.paragraphs:
        text = para.text.strip()
        # Determine style name safely (some paragraphs have style=None)
        style_name = (para.style.name or "") if para.style else ""

        if not text:
            in_tip = False
            continue

        is_heading = "Heading" in style_name or re.match(r"^\d+\.\d+\s", text)

        if is_heading:
            if current_section and current_section["body"]:
                sections.append(current_section)
            current_section = {"heading": text, "body": "", "exam_tip": ""}
            in_tip = False
        elif current_section is not None:
            if "EXAM TIP" in text.upper() and len(text) < 25:
                in_tip = True
                continue
            if in_tip:
                current_section["exam_tip"] += (" " if current_section["exam_tip"] else "") + text
            else:
                sep = "\n" if style_name == "List Paragraph" else " "
                if current_section["body"]:
                    current_section["body"] += sep + text
                else:
                    current_section["body"] = text

    if current_section and current_section["body"]:
        sections.append(current_section)

    if not sections:
        return day_meta

    result = dict(day_meta)
    result["sections"] = sections
    return result


def _parse_questions_docx(path: Path) -> list:
    """Parse the Day 1 practice questions .docx.

    Each question occupies a 3-row table:
      Row 0: ['Q<N>', '<topic> · <difficulty>']
      Row 1: ['Q',    '<question text>\\nA. ...\\nB. ...\\nC. ...\\nD. ...']
      Row 2: ['A',    'Correct Answer: X\\n<explanation>']

    Tables 0 and 1 are a header/TOC and are skipped.
    """
    from docx import Document
    doc = Document(str(path))
    questions = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) != 3:
            continue

        cells_0 = [c.text.strip() for c in rows[0].cells]
        cells_1 = [c.text.strip() for c in rows[1].cells]
        cells_2 = [c.text.strip() for c in rows[2].cells]

        # Row 0: exactly 2 cells, first matches Q<N>
        if len(cells_0) < 2:
            continue
        q_match = re.match(r"Q(\d+)$", cells_0[0])
        if not q_match:
            continue
        q_num = int(q_match.group(1))

        meta_text = cells_0[1]
        difficulty = "basic"
        if "Intermediate" in meta_text:
            difficulty = "intermediate"
        elif "Advanced" in meta_text:
            difficulty = "advanced"

        # Row 1: cell 0 == 'Q', cell 1 == question + options
        if len(cells_1) < 2 or cells_1[0] != "Q":
            continue
        q_body = cells_1[1]

        # Split question text from options (options start at '\nA. ' or 'A. ' on its own line)
        option_split = re.split(r"\n(?=[A-D]\.)", q_body)
        q_clean = option_split[0].strip()
        options = {}
        for part in option_split[1:]:
            m = re.match(r"([A-D])\.\s*(.*)", part, re.DOTALL)
            if m:
                options[m.group(1)] = m.group(2).strip()

        # Row 2: cell 0 == 'A', cell 1 == answer + explanation
        if len(cells_2) < 2 or cells_2[0] != "A":
            continue
        a_body = cells_2[1]
        answer_match = re.search(r"Correct Answer:\s*([A-D])", a_body)
        answer = answer_match.group(1) if answer_match else "A"
        explanation = re.sub(r"Correct Answer:\s*[A-D]\s*", "", a_body).strip()

        if q_clean and len(options) == 4:
            questions.append({
                "id": f"d1_q{q_num}",
                "day": 1,
                "topic": "Subscription Model",
                "difficulty": difficulty,
                "question": q_clean,
                "options": options,
                "answer": answer,
                "explanation": explanation,
            })

    return questions
